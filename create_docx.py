import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    TEAL = RGBColor(0x2A, 0xBA, 0xD2)
    DARK_TITLE = RGBColor(0x0F, 0x17, 0x2A)
    MUTED = RGBColor(0x64, 0x74, 0x8B)
    RED_ACCENT = RGBColor(0xDC, 0x26, 0x26)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("GARY ASHWORTH'S BUSINESS SELLABILITY ASSESSMENT")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = DARK_TITLE
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Gary Ashworth · Master build specification, version 2.\nScoring engine, valuation model, diagnosis logic and final copy.")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = TEAL
    
    doc.add_paragraph()

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = TEAL
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = DARK_TITLE
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = DARK_TITLE
        return p

    def add_callout(text, title=None, border_color="2ABAD2", fill_color="F0FDF4"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
        tcPr.append(shd)
        
        borders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="none"/>
                <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
            </w:tcBorders>
        ''')
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        if title:
            run_t = p.add_run(title + "\n")
            run_t.font.bold = True
            run_t.font.size = Pt(10.5)
            run_t.font.color.rgb = DARK_TITLE
        run_b = p.add_run(text)
        run_b.font.size = Pt(10)
        run_b.font.color.rgb = DARK_TITLE
        doc.add_paragraph()

    # --- 0. HOW THE ENGINE WORKS ---
    add_h1("0. How the Engine Works")
    p = doc.add_paragraph("Read this before building anything else. The scoring in version 1 did not add up to 100, which meant roughly half of all respondents would have been told they owned a trophy business. Everything below assumes the corrected model.")
    
    add_callout(
        "• Fifteen scored questions plus three conditionals. Raw points total 174 at maximum.\n"
        "• Final score = (raw points / 174) x 100, rounded to one decimal, then clamped between 0 and 100.\n"
        "• The recurring revenue bonus is applied to the valuation multiple (+0.5x), not to the score.\n"
        "• Question 16 carries no points. It routes the call to action only.\n"
        "• Estimated revenue and margin figures are backend values and never appear on screen.\n\n"
        "HARD RULES:\n"
        "• Q3 = Declining (Going backwards): Final score capped at 55 (top of Tier 2 at best).\n"
        "• Q5 = Sales drop to zero (They stop): Sets Owner Dependent flag. Suppresses ready to sell language.\n"
        "• Q2 = Negative/breakeven AND Q2B = Structural issues: Suppress pound valuation. Show Broken Economics result.\n"
        "• Q10 = Over 50%: Adds 0.5x to valuation multiple, after size cap.\n"
        "• Q1 = Under 1 million: Multiple capped at 4.0x regardless of score.",
        title="Scoring Rules & Hard Constraints:",
        border_color="2ABAD2",
        fill_color="F0FDFA"
    )

    # --- 1. LANDING PAGE & LEAD CAPTURE ---
    add_h1("1. Landing Page and Lead Capture")
    
    add_callout(
        "Between 70% and 80% of businesses that go to market never sell. Source: Exit Planning Institute.",
        title="⚠️ Banner Statistic:",
        border_color="DC2626",
        fill_color="FEF2F2"
    )

    p = doc.add_paragraph()
    p.add_run("Headline: ").bold = True
    p.add_run("If you put your business on the market next Monday, what would a buyer really pay for it?")
    
    p = doc.add_paragraph()
    p.add_run("Subheadline: ").bold = True
    p.add_run("Answer sixteen questions and I will show you the number, the multiple your business earns today, and the one thing doing the most damage to your price. Includes tailored recommendations")

    p = doc.add_paragraph()
    p.add_run("Duration Badge: ").bold = True
    p.add_run("2–3 minutes to complete")

    p = doc.add_paragraph()
    p.add_run("Credibility Block: ").bold = True
    p.add_run("Forty years of buying, building and selling businesses. Thirty of them, give or take, and a fair few lessons I paid for the hard way.")

    p = doc.add_paragraph()
    p.add_run("Secondary Credibility Block (Results Page / Under the fold): ").bold = True
    p.add_run('"I assumed my tech recruitment business would sell for eight to ten times EBITDA, because that is what the sector had always achieved. By the time I came to sell, the going rate was five. Nobody had told me the market had moved, and that gap cost me millions. This diagnostic exists so it does not happen to you."')

    add_h2("What You Get at the End")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Your number: ").bold = True
    p.add_run("What a buyer would likely pay today, and the multiple your business is earning.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Your biggest value killer: ").bold = True
    p.add_run("The single risk taking the most money off your price, named and costed.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Your gap: ").bold = True
    p.add_run("The difference in pounds between what you would get today and the ceiling for a business your size, with the order I would fix things in.")

    add_h2("Form Fields")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Full name, required")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Work email, required")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Company name, optional")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Button copy: ").bold = True
    p.add_run("Start the diagnostic")

    # --- 2. QUESTIONS AND LOGIC ---
    add_h1("2. Questions and Logic (Copy is Final)")

    def render_q(q_num, title, subtitle, choices, note=None):
        add_h3(f"{q_num}. {title}")
        p_s = doc.add_paragraph()
        r_s = p_s.add_run(f"Context: {subtitle}")
        r_s.font.italic = True
        r_s.font.color.rgb = MUTED
        for c, pts in choices:
            p_c = doc.add_paragraph(style='List Bullet')
            p_c.add_run(f"{c} ").bold = True
            p_c.add_run(f"({pts})")
        if note:
            add_callout(note, title="Build Note:", border_color="2ABAD2", fill_color="F0F9FF")

    add_h2("Section 1: The money")
    render_q(
        "Q1", "What was your turnover over the last twelve months?",
        "Revenue decides which buyers will look at you at all, and it sets the ceiling on what any of them will pay.",
        [("Under 1 million", "5 pts, backend estimate 650k"), ("1 million to 3 million", "10 pts, backend estimate 2.0m"), ("3 million to 10 million", "15 pts, backend estimate 6.0m"), ("Over 10 million", "20 pts, backend estimate 15.0m")]
    )

    render_q(
        "Q2", "What is your net profit margin before your own salary, interest and tax?",
        "This is the number buyers price off. Turnover is what you tell people at dinner, margin is what somebody buys.",
        [("Negative or breakeven", "0 pts, triggers Q2B, backend estimate 3%"), ("1% to 10%", "5 pts, backend estimate 7.5%"), ("11% to 20%", "10 pts, backend estimate 15%"), ("Over 20%", "15 pts, backend estimate 25%")]
    )

    render_q(
        "Q2B (conditional)", "Why is the business not making money?",
        "Losing money on purpose whilst you grow is a very different conversation to losing money because something is broken.",
        [("Deliberate reinvestment to grow faster", "+3 pts"), ("Temporary market conditions", "+1 pt"), ("Pricing is wrong", "0 pts"), ("Structural problems in the business", "-5 pts, suppresses pound valuation")]
    )

    render_q(
        "Q3", "How is turnover moving year on year?",
        "Buyers are paying for the next three years of profit, and a falling line tells them the next three will be worse than the last three.",
        [("Going backwards", "0 pts, caps final score at 55"), ("Flat", "5 pts"), ("Growing 1% to 15%", "10 pts"), ("Growing more than 15%", "15 pts")]
    )

    add_h2("Section 2: How much of it is you")
    render_q(
        "Q4", "How many hours a week do you spend on the day to day work?",
        "The more of the daily running that goes through you, the harder it is for anyone else to take over.",
        [("More than 40", "0 pts"), ("20 to 40", "3 pts"), ("5 to 20", "7 pts"), ("Under 5", "10 pts")]
    )

    render_q(
        "Q5", "What happens to sales if you disappear for three months?",
        "This is the question that separates a business from a well paid job.",
        [("They stop", "0 pts, sets Owner Dependent flag"), ("They halve", "3 pts"), ("They hold steady", "7 pts"), ("They keep growing, because the team runs it", "10 pts")]
    )

    render_q(
        "Q6", "Do you have a number two who can run the place without you?",
        "A proper second in command moves your final price more than almost anything else on this list.",
        [("No", "0 pts"), ("Yes, but they are new", "3 pts"), ("Yes, with some experience", "7 pts"), ("Yes, and they have a long track record", "12 pts")]
    )

    add_h2("Section 3: Where it could all go wrong")
    render_q(
        "Q7", "What share of turnover comes from your largest customer?",
        "One big customer is the thing that most often kills a deal at the last minute, usually the week the buyer's bank gets involved.",
        [("Under 5%", "10 pts"), ("5% to 15%", "7 pts"), ("16% to 30%", "3 pts"), ("Over 30%", "-5 pts, triggers Q7B")]
    )

    render_q(
        "Q7B (conditional)", "Is that customer on a contract?",
        "A signed long term agreement changes how a buyer views the same concentration.",
        [("Long term binding contract", "+3 pts"), ("Rolling month to month", "0 pts"), ("Pays upfront", "+1 pt"), ("Pays late, every time", "-3 pts")]
    )

    render_q(
        "Q8", "How easily could you replace your main supplier?",
        "A supplier you cannot replace is somebody else's hand on your margin.",
        [("One supplier, no realistic alternative", "0 pts, triggers Q8B"), ("Very difficult, it would cost us time and margin", "2 pts"), ("Doable with some disruption", "6 pts"), ("Easy, several suppliers compete for our business", "10 pts")]
    )

    render_q(
        "Q8B (conditional)", "Is that single source something you own or control, such as an exclusive licence, a patent or an agreement nobody else can get?",
        "Owning the only route to market is an asset. Being hostage to somebody else's is a risk.",
        [("Yes, we own or control it", "+8 pts"), ("No, we are dependent on them", "0 pts")]
    )

    render_q(
        "Q9", "What happens if your best salesperson resigns tomorrow?",
        "This tells a buyer whether the sales system belongs to the business or lives in one person's head.",
        [("We would lose a serious chunk of revenue", "0 pts"), ("A big drop for several months whilst we recover", "3 pts"), ("A dip, then back to normal", "7 pts"), ("Barely noticed, the leads and the process belong to the business", "10 pts")]
    )

    add_h2("Section 4: Quality of the earnings")
    render_q(
        "Q10", "How much of your revenue is on subscription or contract?",
        "Contracted income is the closest thing to guaranteed money a buyer can see, and they pay a premium for it.",
        [("None, every sale starts from scratch", "0 pts"), ("Under 20%", "4 pts"), ("21% to 50%", "8 pts"), ("Over 50%", "12 pts, adds 0.5x to multiple")]
    )

    render_q(
        "Q11", "How would your customers react to a 10% price rise tomorrow?",
        "If you can put prices up without losing people, you have something they cannot get elsewhere.",
        [("We would lose a lot of them", "0 pts"), ("Plenty of complaints and some would go", "3 pts"), ("Some grumbling, most would stay", "7 pts"), ("Nobody would blink", "10 pts")]
    )

    render_q(
        "Q12", "How do your gross margins compare with others in your sector?",
        "Margin above your sector average tells a buyer you have pricing power and room to absorb a bad year.",
        [("Well below average", "0 pts"), ("Slightly below", "3 pts"), ("About the same", "7 pts"), ("Well above average", "10 pts")]
    )

    render_q(
        "Q13", "When do your customers pay you?",
        "Getting paid before you deliver means a buyer needs less cash to run the place, and that shows up in the price.",
        [("60 days or more", "0 pts"), ("30 days", "3 pts"), ("On delivery", "6 pts"), ("Upfront, before we do the work", "10 pts")]
    )

    render_q(
        "Q14", "How likely are your customers to recommend you?",
        "Customers who bring you more customers keep your acquisition cost down, and a buyer will notice.",
        [("Very unlikely", "0 pts"), ("Unlikely", "3 pts"), ("Neither one way nor the other", "6 pts"), ("Very likely, they do it already", "10 pts")]
    )

    render_q(
        "Q15", "How are your accounts kept?",
        "Messy books give a buyer an excuse to chip away at your price during due diligence, and they will take it.",
        [("Shoebox, or whatever the accountant can piece together", "0 pts"), ("Xero or QuickBooks, kept up to date", "4 pts"), ("Accountant prepared and filed each year", "7 pts"), ("Fully audited by an independent firm", "10 pts")]
    )

    add_h2("Section 5: You")
    render_q(
        "Q16", "When do you want to be out?",
        "This one carries no points. It changes what I would tell you to do first.",
        [("Inside twelve months", "0 pts, routes to call"), ("One to three years", "0 pts, routes by tier"), ("Three to five years", "0 pts, routes by tier"), ("No fixed plan, I just want to know where I stand", "0 pts, routes by tier")]
    )

    # --- 3. VALUATION ENGINE ---
    add_h1("3. Valuation Engine")
    add_callout(
        "1. Earnings: Adjusted EBITDA = revenue estimate from Q1 x margin estimate from Q2\n"
        "2. Score: (raw points / 174) x 100, clamped 0 to 100, then apply Q3 cap of 55 if it fires\n"
        "3. Base multiple: 0 to 39.9 = 2.0x, 40 to 61.9 = 3.8x, 62 to 79.9 = 5.8x, 80 to 100 = 8.5x\n"
        "4. Recurring uplift: Add 0.5x if Q10 = over 50%\n"
        "5. Size cap: Under 1m capped at 4.0x, 1m to 3m at 5.5x, 3m to 10m at 7.0x, over 10m at 8.5x\n"
        "6. Today's value: Adjusted EBITDA x final multiple\n"
        "7. Ceiling: Adjusted EBITDA x the size cap for their revenue band\n"
        "8. The gap: Ceiling minus today's value",
        title="8-Step Calculation Workflow:",
        border_color="2ABAD2",
        fill_color="F0FDFA"
    )

    # --- 4. NUMBER ONE VALUE KILLER ---
    add_h1("4. The Number One Value Killer (Ordered Priority)")
    
    def render_vk(num, title, trig, cost, diag, fix):
        add_h2(f"{num}. {title}")
        p = doc.add_paragraph()
        p.add_run("Trigger: ").bold = True
        p.add_run(trig)
        p = doc.add_paragraph()
        p.add_run("What it costs: ").bold = True
        p.add_run(cost)
        p = doc.add_paragraph()
        p.add_run("Diagnosis: ").bold = True
        p.add_run(diag)
        p = doc.add_paragraph()
        p.add_run("What to do about it: ").bold = True
        p.add_run(fix)
        doc.add_paragraph()

    render_vk("1", "You are the business", "Q5 = sales stop", "Halves the multiple and puts you on an earn out for three to five years", "Sales stop if you step away for three months, which means the business is you, and nobody can buy you. A trade buyer will walk. A private equity buyer will make you an offer with most of the money tied to an earn out that keeps you working for another three to five years to collect it.", "Write down how the work gets done, one process at a time. Hand your ten biggest accounts to somebody else over the next six months. Get one person owning delivery and one owning the pipeline, and stay out of both.")
    render_vk("2", "One customer holds the keys", "Q7 = over 30%", "Two to three turns off the multiple, and most lenders will not fund the deal at all", "More than thirty percent of your turnover sits in one account. Ask yourself what happens to a buyer if that client leaves the month after completion, because the buyer's bank will ask exactly that, and most lenders refuse to fund a deal with more than a quarter of revenue in one place.", "Get that client onto a multi year agreement with proper notice periods, and put a serious push behind new business until they sit under fifteen percent. Both jobs, not one or the other.")
    render_vk("3", "The line is going backwards", "Q3 = declining", "Score capped at 55 and a turnaround discount on top", "Turnover is going backwards. Buyers price the next three years, and a falling line tells them the next three are worse than the last three. They will treat you as a turnaround and pay you like one, if they bid at all. Fix the direction before you go anywhere near a broker.", "Cut the lines that make no money, put prices up where the market lets you, and get two consecutive quarters of growth on the board before you speak to a single buyer.")
    render_vk("4", "The numbers do not work", "Q2 = negative/breakeven and Q2B = structural problems", "No earnings multiple applies. Asset value only", "The business is not making money and you have told me the reason is structural rather than deliberate. There is no earnings multiple to apply to a business that does not earn, so the only offer on the table would be for the assets. That is a business problem before it is an exit problem.", "Work out which customers and which product lines make you money, and be willing to lose the ones that do not. Get to a positive number for two quarters before anything else on this list matters.")
    render_vk("5", "There is nobody behind you", "Q6 = no number two", "One to two turns off the multiple", "There is nobody who can run the place without you. That one hire usually moves the final price more than a year of extra sales does, and most owners leave it far too late because they cannot see how to afford it. The maths says you cannot afford not to.", "Hire or promote somebody into the number two seat this quarter, give them real authority rather than the title, and let them make a few expensive mistakes whilst you are still there to catch them.")
    render_vk("6", "Your books will not survive diligence", "Q15 = shoebox", "Five to fifteen percent knocked off during diligence, plus months of delay", "Your records will not stand up to a proper look. Every number a buyer cannot trace becomes a reason to pay you less, and buyers are very good at finding them. This is the cheapest thing on the list to fix and the one most owners put off.", "Get twenty four months of clean, reconciled accounts prepared properly, separate anything personal that is running through the business, and have your accountant produce a set that a stranger could follow.")
    render_vk("7", "Every pound has to be won again", "Q10 = none", "One to two turns off the multiple", "None of your revenue renews on its own, so every year starts at zero. Buyers pay a great deal more for income that arrives whether or not anybody sells anything, because it is the difference between buying a business and buying a treadmill.", "Find the part of what you do that customers need every month, price it separately, and put your best accounts onto twelve month agreements.")
    render_vk("8", "You cannot move your prices", "Q11 = we would lose a lot of them", "Signals a commodity position, which caps the multiple", "A ten percent price rise would cost you customers, which tells a buyer you are competing on price against people who can undercut you. That caps what anybody will pay, because they can see the margin has nowhere to go.", "Work out what you do that the cheap option does not, put it in front of customers, and test a price rise on your ten best accounts before you touch the rest.")

    # --- 5. RESULT TIERS ---
    add_h1("5. Result Tiers")

    def render_t(tier_title, score_range, mult, hline, desc, roadmap):
        add_h2(f"{tier_title} (Score {score_range}, {mult})")
        p = doc.add_paragraph()
        p.add_run("Headline: ").bold = True
        p.add_run(hline)
        p = doc.add_paragraph()
        p.add_run("Description: ").bold = True
        p.add_run(desc)
        add_h3("The Roadmap:")
        for idx, step in enumerate(roadmap, 1):
            p_s = doc.add_paragraph(style='List Bullet')
            p_s.add_run(f"Step {idx}: ").bold = True
            p_s.add_run(step)
        doc.add_paragraph()

    render_t("Tier 1. You own a job", "0 to 39.9", "around 2.0x", "Right now a buyer sees a demanding job with your name on it", "Buyers want a business that runs whether or not you turn up. At this score they would be buying themselves a job, and most of them already have one. That is a fixable problem, and it is the one owners fix too late. The work below is unglamorous, it takes six to twelve months, and it moves your number more than any sales push will.", [
        "Write down how the core work gets done, one process at a time, starting with whatever only you know.",
        "Promote or hire somebody to run delivery and the team day to day.",
        "Get leads coming in through a system rather than through your phone.",
        "Move your key client relationships onto your people over six months, and stay off the calls."
    ])

    render_t("Tier 2. Sellable, and the price gets chipped", "40 to 61.9", "around 3.8x", "You could sell this, and you would watch a buyer take money off you for six weeks", "There is a real business here and somebody would buy it. You would also spend the back end of the process watching a buyer knock the price down for risks you could have dealt with beforehand. Every item below is a discount you are currently handing over.", [
        "Bring your largest customer under fifteen percent of turnover.",
        "Get your top accounts onto one to three year agreements with clear renewal terms.",
        "Sort out contracts, notice periods and retention for the people a buyer would worry about losing.",
        "Clean up two years of accounts now, whilst you have time, rather than during diligence."
    ])

    render_t("Tier 3. Buyers will compete for this", "62 to 79.9", "around 5.8x", "You have built something clean, and the job now is running a proper process", "Margins hold up, somebody else can run it, and the revenue does not hang on one client or one person. Businesses like yours get bid for. What loses money at this stage is selling to the first buyer who calls, because a single bidder sets the price and knows it.", [
        "Get a quality of earnings review done before a buyer does one to you.",
        "Convert whatever you can from one off work to contracted income, because it is worth more than the revenue it replaces.",
        "Write the three year growth story, with the channels you have not touched yet and what they are worth.",
        "Interview three advisers and pick the one who will get more than one buyer to the table."
    ])

    render_t("Tier 4. Top of the market", "80 to 100", "around 8.5x", "Very few owners get here", "The management runs it, the income renews, the books stand up and the contracts hold. Your job now is making sure you do not leave money on the table by taking the first approach that lands in your inbox, which is how most owners at your level end up underpaid.", [
        "Run a proper private auction through an adviser who works in your sector.",
        "Look hard at trade buyers who gain something from owning you, because they pay above the market rate.",
        "Get the tax structuring done early, and by somebody who does this for a living.",
        "Lock in your key people with retention packages before anybody knows the business is for sale."
    ])

    render_t("Tier 4 variant (Under 1m turnover)", "80 to 100", "4.0x size cap", "As good as it gets at your size", "The fundamentals here are as strong as I see them. Scale is the only thing holding the multiple down, because the buyers who pay the highest multiples do not look below a few million in turnover. Growth is worth more to your final number now than any further tidying up, and you already have the machine to grow with.", [
        "Put everything into growth for the next two years, since the fundamentals are already right.",
        "Keep the recurring revenue percentage climbing whilst you scale.",
        "Revisit this diagnostic when you cross three million, because the ceiling on your multiple moves at that point."
    ])

    # --- 6. WHERE EACH RESULT LEADS ---
    add_h1("6. Where Each Result Leads (Call to Action Routing)")
    
    add_callout(
        "ROUTING RULES:\n"
        "• Q16 = inside twelve months (any tier) OR Tier 3/4 ➔ Route to Call\n"
        "• Tier 1 or Tier 2 (exit > 12 months out) ➔ Route to Accelerator\n"
        "• Broken Economics ➔ Route to Accelerator\n\n"
        "CTA FOR THE CALL:\n"
        "Headline: Thirty minutes, your numbers, no pitch.\n"
        "Body: If you want the unvarnished version, book half an hour with me and bring your figures. We will go through what is costing you the most, what it is worth in pounds, and the order I would fix it in if this were my business. No sales pitch and no generic advice, because neither of us has time for it.\n"
        "Button: Book your thirty minute review with Gary\n\n"
        "CTA FOR THE ACCELERATOR:\n"
        "Headline: The work comes before the exit.\n"
        "Body: Your score says the job right now is building the thing a buyer would want, and that is months of work rather than a phone call. The Accelerator is where I take owners through it, in order, using the same checklist I run over any business before I buy it.\n"
        "Button: Join the Business Accelerator",
        title="Conversion Routing Logic:",
        border_color="2ABAD2",
        fill_color="F8FAFC"
    )

    output_path = "/Users/jamesmacintosh/.gemini/antigravity/scratch/Gary - Business Sellability/Business_Sellability_Assessment_Master_Architecture.docx"
    doc.save(output_path)
    print(f"Docx created successfully at {output_path}")

if __name__ == "__main__":
    create_document()
